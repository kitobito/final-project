# backend/main.py

import os
import tempfile
import json
from dotenv import load_dotenv

import db
from db import engine, Base

# IMPORT MODELS so that SG create_all() sees them
import models

load_dotenv()
GROQ_KEY = os.getenv("GROQ_API_KEY")
if not GROQ_KEY:
    raise RuntimeError("❌ Please set GROQ_API_KEY in backend/.env")

# Groq HTTP settings
GROQ_URL = "https://api.groq.com/openai/v1"
GROQ_HEADERS = {
    "Authorization": f"Bearer {GROQ_KEY}",
    "Content-Type": "application/json"
}

# Create tables (users, conversations, chat_messages)
Base.metadata.create_all(bind=engine)

from fastapi import FastAPI, UploadFile, File, Query, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import List, Dict

import pdfplumber
import requests
from docx import Document

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
def root():
    return {"message": "Backend is running"}

# ——————————————————————————————————————————————
# 1) File Upload & Plain Summarization (unchanged)
@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    text = ""
    if file.filename.lower().endswith(".pdf"):
        with pdfplumber.open(file.file) as pdf:
            for page in pdf.pages:
                p = page.extract_text() or ""
                text += p + "\n"
    else:
        data = await file.read()
        text = data.decode("utf-8", errors="ignore")

    snippet = text[:6000]
    payload = {
        "model": "llama3-70b-8192",
        "messages": [
            {"role": "system", "content": "You are an expert research assistant. Summarize this text."},
            {"role": "user", "content": snippet}
        ],
        "temperature": 0.3
    }
    r = requests.post(f"{GROQ_URL}/chat/completions", headers=GROQ_HEADERS, json=payload)
    r.raise_for_status()
    data = r.json()
    summary = data["choices"][0]["message"]["content"]
    return {"summary": summary}

# ——————————————————————————————————————————————
# 2) Web Search via DuckDuckGo (unchanged)
@app.get("/search")
def web_search(q: str = Query(..., description="Search query")):
    r = requests.get(
        "https://api.duckduckgo.com/",
        params={"q": q, "format": "json", "no_html": 1}
    )
    r.raise_for_status()
    dd = r.json()
    results = []
    if dd.get("AbstractText"):
        results.append({
            "title": "Summary",
            "snippet": dd["AbstractText"],
            "url": dd.get("AbstractURL", "")
        })
    for topic in dd.get("RelatedTopics", [])[:5]:
        if isinstance(topic, dict) and topic.get("Text") and topic.get("FirstURL"):
            results.append({
                "title": topic["Text"],
                "snippet": topic["Text"],
                "url": topic["FirstURL"]
            })
    return {"results": results}

# ——————————————————————————————————————————————
# 3) Structured Summary Endpoint (unchanged)
@app.post("/summarize")
def summarize_structured(
    text: str = Body(..., embed=True),
    fmt: str = Body(..., embed=True)  # "bullet" or "json"
):
    system_prompt = (
        f"You are a research assistant. "
        f"Please summarize the following text in *{fmt}* format:\n\n"
        f"{text}"
    )
    payload = {
        "model": "llama3-70b-8192",
        "messages": [{"role": "system", "content": system_prompt}],
        "temperature": 0.3
    }
    r = requests.post(f"{GROQ_URL}/chat/completions", headers=GROQ_HEADERS, json=payload)
    r.raise_for_status()
    data = r.json()
    return {"summary": data["choices"][0]["message"]["content"]}

# ——————————————————————————————————————————————
# 4) Analysis Endpoint (unchanged)
@app.post("/analyze")
def analyze_text(text: str = Body(..., embed=True)):
    system_prompt = (
        "You are a data extractor. "
        "Extract key findings from the text below and output a JSON array ONLY. "
        "Each entry should be an object with two keys: metric and value. "
        "Do NOT include any explanatory text or Markdown—only raw JSON.\n\n"
        f"{text}"
    )

    payload = {
        "model": "llama3-70b-8192",
        "messages": [{"role": "system", "content": system_prompt}],
        "temperature": 0.0
    }

    r = requests.post(f"{GROQ_URL}/chat/completions", headers=GROQ_HEADERS, json=payload)
    r.raise_for_status()
    data = r.json()
    content = data["choices"][0]["message"]["content"].strip()

    try:
        parsed = json.loads(content)
        return {"data": parsed}
    except Exception:
        return {"data": [], "raw": content}

# ——————————————————————————————————————————————
# 5) Export to .docx (unchanged)
@app.post("/export")
def export_report(
    summary: str = Body(..., embed=True),
    structured: str = Body("", embed=True)
):
    doc = Document()
    doc.add_heading("SynthesisTalk Report", level=1)
    doc.add_paragraph("=== Plain Summary ===")
    doc.add_paragraph(summary)
    if structured:
        doc.add_paragraph("\n=== Structured Summary ===")
        doc.add_paragraph(structured)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return FileResponse(
        tmp.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="report.docx"
    )

# ——————————————————————————————————————————————
# 6) Include authentication, conversation, and chat routers
import auth
import conversation_router
import chat_router

from auth import router as auth_router
from conversation_router import router as conv_router
from chat_router import router as chat_router

app.include_router(auth_router)
app.include_router(conv_router)
app.include_router(chat_router)
